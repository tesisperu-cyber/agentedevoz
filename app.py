import streamlit as st
import time
import io
import re
from datetime import datetime

# ─── Page config ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Marco Teórico IA",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Custom CSS ─────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&family=IBM+Plex+Sans:wght@300;400;500&family=IBM+Plex+Mono:wght@400;500&display=swap');

/* Base */
html, body, [class*="css"] {
    font-family: 'IBM Plex Sans', sans-serif;
}

/* App background */
.stApp {
    background: #f7f5f0;
}

/* Sidebar */
[data-testid="stSidebar"] {
    background: #1a1b1e !important;
    border-right: 1px solid #2e2f33;
}
[data-testid="stSidebar"] * {
    color: #d4cfc4 !important;
}
[data-testid="stSidebar"] .stSelectbox label,
[data-testid="stSidebar"] .stTextInput label,
[data-testid="stSidebar"] .stTextArea label {
    color: #8a8478 !important;
    font-family: 'IBM Plex Mono', monospace !important;
    font-size: 0.7rem !important;
    text-transform: uppercase;
    letter-spacing: 0.1em;
}
[data-testid="stSidebar"] input,
[data-testid="stSidebar"] textarea,
[data-testid="stSidebar"] select {
    background: #0f1012 !important;
    border: 1px solid #2e2f33 !important;
    color: #d4cfc4 !important;
    font-family: 'IBM Plex Mono', monospace !important;
    font-size: 0.82rem !important;
}
[data-testid="stSidebar"] .stSelectbox > div > div {
    background: #0f1012 !important;
    border: 1px solid #2e2f33 !important;
    color: #d4cfc4 !important;
}

/* Header banner */
.main-header {
    background: linear-gradient(135deg, #1a1b1e 0%, #242529 100%);
    border-radius: 12px;
    padding: 2.2rem 2.5rem;
    margin-bottom: 1.8rem;
    border-left: 5px solid #c8a96e;
    box-shadow: 0 4px 24px rgba(0,0,0,0.12);
}
.main-header h1 {
    font-family: 'Playfair Display', serif !important;
    color: #c8a96e !important;
    font-size: 2.1rem !important;
    margin: 0 !important;
    letter-spacing: -0.02em;
}
.main-header p {
    color: #6b6560 !important;
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.75rem;
    text-transform: uppercase;
    letter-spacing: 0.12em;
    margin: 0.4rem 0 0 0;
}

/* Section cards */
.section-card {
    background: white;
    border-radius: 10px;
    padding: 1.6rem 1.8rem;
    margin-bottom: 1.2rem;
    border: 1px solid #e8e3da;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}
.section-title {
    font-family: 'Playfair Display', serif;
    color: #1a1b1e;
    font-size: 1.15rem;
    font-weight: 700;
    border-bottom: 2px solid #c8a96e;
    padding-bottom: 0.5rem;
    margin-bottom: 1rem;
}
.section-content {
    font-size: 0.87rem;
    line-height: 1.7;
    color: #3a3732;
    white-space: pre-wrap;
}

/* Progress */
.progress-item {
    display: flex;
    align-items: center;
    gap: 0.6rem;
    padding: 0.4rem 0;
    font-size: 0.83rem;
    color: #3a3732;
}
.dot-done { color: #2ecc71; font-size: 1rem; }
.dot-working { color: #c8a96e; font-size: 1rem; }
.dot-pending { color: #ccc; font-size: 1rem; }

/* Metric boxes */
.metric-row {
    display: flex;
    gap: 1rem;
    margin-bottom: 1.2rem;
}
.metric-box {
    flex: 1;
    background: #1a1b1e;
    border-radius: 8px;
    padding: 1rem 1.2rem;
    text-align: center;
    border: 1px solid #2e2f33;
}
.metric-box .num {
    font-family: 'Playfair Display', serif;
    font-size: 1.8rem;
    color: #c8a96e;
    font-weight: 700;
    line-height: 1;
}
.metric-box .lbl {
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.65rem;
    color: #6b6560;
    text-transform: uppercase;
    letter-spacing: 0.1em;
    margin-top: 0.2rem;
}

/* Generate button */
.stButton > button {
    background: linear-gradient(135deg, #c8a96e, #b8924a) !important;
    color: #1a1b1e !important;
    font-family: 'IBM Plex Mono', monospace !important;
    font-weight: 600 !important;
    font-size: 0.9rem !important;
    border: none !important;
    border-radius: 8px !important;
    padding: 0.7rem 2rem !important;
    width: 100% !important;
    letter-spacing: 0.05em;
    text-transform: uppercase;
    transition: all 0.2s !important;
}
.stButton > button:hover {
    background: linear-gradient(135deg, #d4b87a, #c8a96e) !important;
    box-shadow: 0 4px 16px rgba(200,169,110,0.35) !important;
}

/* Download button */
.stDownloadButton > button {
    background: #1a1b1e !important;
    color: #c8a96e !important;
    font-family: 'IBM Plex Mono', monospace !important;
    font-weight: 500 !important;
    border: 1px solid #c8a96e !important;
    border-radius: 8px !important;
    width: 100% !important;
}

/* Sidebar brand */
.sidebar-brand {
    font-family: 'Playfair Display', serif;
    font-size: 1.1rem;
    color: #c8a96e !important;
    padding: 0.5rem 0 1rem 0;
    border-bottom: 1px solid #2e2f33;
    margin-bottom: 1rem;
}
.sidebar-tag {
    font-family: 'IBM Plex Mono', monospace;
    font-size: 0.62rem;
    color: #4a4843 !important;
    text-transform: uppercase;
    letter-spacing: 0.1em;
}

/* Alert / info box */
.info-box {
    background: #fffbf2;
    border: 1px solid #e8d5a3;
    border-left: 4px solid #c8a96e;
    border-radius: 6px;
    padding: 0.8rem 1rem;
    font-size: 0.82rem;
    color: #5a4e38;
    margin-bottom: 1rem;
}

/* Hide streamlit elements */
#MainMenu, footer, header { visibility: hidden; }
.block-container { padding-top: 1.5rem !important; }
</style>
""", unsafe_allow_html=True)


# ─── Helpers ────────────────────────────────────────────────────────────────

def get_groq_client(api_key: str):
    from groq import Groq
    return Groq(api_key=api_key)


def llamar_groq(client, prompt: str, max_tokens: int = 2800) -> str:
    sistema = """Eres un experto en metodología de investigación científica con doctorado.
Redactas marcos teóricos académicos rigurosos en español formal.
Incluyes citas de autores en español, inglés, portugués y chino mandarín (traducidos al español, indicando [Traducido del chino mandarín]).
Cuando citas autores chinos, escribes primero el nombre en español/pinyin y luego la indicación.
Usas lenguaje formal, preciso y científico. Años de publicación preferentemente 2018-2026."""

    resp = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": sistema},
            {"role": "user", "content": prompt},
        ],
        temperature=0.72,
        max_tokens=max_tokens,
    )
    return resp.choices[0].message.content.strip()


# ─── Prompts ────────────────────────────────────────────────────────────────

def prompt_definiciones(datos, variable):
    return f"""Para la variable "{variable}" en el contexto de una investigación sobre:
"{datos['tema']}"
(Tipo: {datos['tipo_estudio']}, Área: {datos['area']}, Norma: {datos['norma']})

Proporciona DEFINICIONES CONCEPTUALES redactadas por al menos 5 autores diferentes:
- 2 autores hispanohablantes
- 1 autor anglófono (con traducción al español)
- 1 autor lusófono (con traducción al español)
- 1 autor chino (con nombre en pinyin y traducción al español, indicar [Traducido del chino mandarín])

Para cada definición incluye análisis conceptual en párrafo independiente y cita en formato {datos['norma']}.
Años: 2018-2026. Redacta con rigor académico."""


def prompt_teorias(datos, variable):
    dim = "categorías" if datos['tipo_estudio'] == "Cualitativo" else "dimensiones"
    return f"""Para la variable "{variable}" — investigación sobre: "{datos['tema']}"
(Tipo: {datos['tipo_estudio']}, Área: {datos['area']}, Norma: {datos['norma']})

Desarrolla en formato académico:

1. TEORÍAS / MODELOS TEÓRICOS (mínimo 3):
   Para cada uno: nombre, autor(es) con año, descripción detallada, aplicación al estudio, cita en {datos['norma']}.

2. {dim.upper()} DE LA VARIABLE (mínimo 4):
   Para cada {dim[:-1]}: nombre, definición con cita, indicadores o subcategorías.

Autores en español, inglés, portugués y chino (traducido). Años 2018-2026."""


def prompt_glosario(datos):
    return f"""Investigación: "{datos['tema']}"
Variables: {', '.join(datos['variables'])}
Área: {datos['area']}, Tipo: {datos['tipo_estudio']}, Norma: {datos['norma']}

Elabora un GLOSARIO DE TÉRMINOS BÁSICOS con exactamente 12 términos clave.
Para cada término:
**TÉRMINO:** Nombre en mayúsculas
Definición académica rigurosa (3-4 líneas).
(Autor, año, en formato {datos['norma']})

Incluye términos técnicos del área, conceptos metodológicos y constructos teóricos.
Autores en español, inglés, portugués y chino mandarín (traducido). Años 2018-2026."""


def prompt_antec_internacionales(datos):
    return f"""Genera exactamente 10 ANTECEDENTES INTERNACIONALES para la investigación sobre:
"{datos['tema']}"
Variables: {', '.join(datos['variables'])}
Área: {datos['area']}, Tipo: {datos['tipo_estudio']}, Norma: {datos['norma']}

REQUISITOS:
- Solo estudios publicados entre 2021 y 2026
- Países distintos a {datos['pais']}
- Incluir al menos: 2 estudios de habla hispana (distintos a {datos['pais']}), 3 en inglés, 1 en portugués, 1 de Asia (China/Japón/Corea)

Para CADA antecedente (numera del 1 al 10):
**Autor(es) (Año).** "Título completo del estudio."
*País | Institución/Universidad*
**Resumen:** [párrafo de 5-7 líneas con: objetivo general, metodología, muestra/participantes, principales resultados y conclusiones]
**Relevancia:** [2 líneas indicando cómo aporta a la presente investigación]
**Cita {datos['norma']}:** [referencia completa]

Si el estudio original es en otro idioma, indicar [Traducido del inglés/portugués/chino]."""


def prompt_antec_nacionales(datos):
    return f"""Genera exactamente 8 ANTECEDENTES NACIONALES de {datos['pais']} para la investigación sobre:
"{datos['tema']}"
Variables: {', '.join(datos['variables'])}
Área: {datos['area']}, Tipo: {datos['tipo_estudio']}, Norma: {datos['norma']}

REQUISITOS:
- Solo estudios publicados entre 2021 y 2026
- Estudios realizados en {datos['pais']} (universidades o centros de investigación del país)

Para CADA antecedente (numera del 1 al 8):
**Autor(es) (Año).** "Título completo del estudio."
*Universidad / Institución — {datos['pais']}*
**Resumen:** [párrafo de 5-7 líneas con: objetivo general, metodología, muestra/participantes, principales resultados y conclusiones]
**Relevancia:** [2 líneas indicando cómo aporta a la presente investigación]
**Cita {datos['norma']}:** [referencia completa]"""


# ─── DOCX Builder ───────────────────────────────────────────────────────────

def build_docx(datos: dict, secciones: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Márgenes
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2.5)

    # Fuente base
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    # ── PORTADA
    for _ in range(4):
        doc.add_paragraph("")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MARCO TEÓRICO")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x35, 0x64)

    doc.add_paragraph("")
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(datos['tema'])
    r2.bold = True; r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("")
    info_lines = [
        f"Tipo de documento: {datos['tipo_doc']}",
        f"Tipo de estudio: {datos['tipo_estudio']}",
        f"Área de conocimiento: {datos['area']}",
        f"País (antecedentes nacionales): {datos['pais']}",
        f"Variables: {', '.join(datos['variables'])}",
        f"Norma de citación: {datos['norma']}",
        f"Generado con IA: {datetime.now().strftime('%d de %B de %Y')}",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line).font.size = Pt(11)

    doc.add_page_break()

    # ── Helper: heading
    def h(text, level):
        p = doc.add_heading(text, level=level)
        run = p.runs[0]
        run.bold = True
        colors = {1: RGBColor(0x1F,0x35,0x64), 2: RGBColor(0x2E,0x74,0xB5), 3: RGBColor(0x1F,0x60,0x80)}
        sizes  = {1: 16, 2: 13, 3: 12}
        run.font.size = Pt(sizes.get(level, 12))
        run.font.color.rgb = colors.get(level, RGBColor(0,0,0))
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(6)
        return p

    # ── Helper: add rich text
    def add_text(text):
        for line in text.split('\n'):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            # Detectar encabezados markdown
            if stripped.startswith("### "):
                h(stripped[4:], 3); continue
            if stripped.startswith("## "):
                h(stripped[3:], 2); continue
            if stripped.startswith("# "):
                h(stripped[2:], 3); continue

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.first_line_indent = Inches(0.35)

            # Negritas **texto**
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # ── CONTENIDO
    h("CAPÍTULO II: MARCO TEÓRICO", 1)
    doc.add_paragraph("")
    h("2.1. Bases Teóricas", 2)

    for i, var in enumerate(datos['variables'], 1):
        h(f"2.1.{i}. Variable: {var}", 3)
        h(f"2.1.{i}.1. Definiciones Conceptuales", 3)
        add_text(secciones[f"def_{var}"])
        doc.add_paragraph("")
        h(f"2.1.{i}.2. Teorías, Modelos Teóricos y Dimensiones", 3)
        add_text(secciones[f"teo_{var}"])
        doc.add_paragraph("")

    doc.add_page_break()
    h("2.2. Glosario de Términos Básicos", 2)
    add_text(secciones["glosario"])

    doc.add_page_break()
    h("2.3. Antecedentes de la Investigación", 2)
    h("2.3.1. Antecedentes Internacionales", 3)
    add_text(secciones["antec_int"])
    doc.add_paragraph("")
    h(f"2.3.2. Antecedentes Nacionales ({datos['pais']})", 3)
    add_text(secciones["antec_nac"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── SIDEBAR ────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown('<div class="sidebar-brand">🎓 Marco Teórico IA<br><span class="sidebar-tag">llama-3.3-70b · Groq</span></div>', unsafe_allow_html=True)

    api_key = st.text_input(
        "API KEY DE GROQ",
        type="password",
        placeholder="gsk_...",
        help="Obtén tu clave gratuita en console.groq.com"
    )
    st.markdown('<div class="sidebar-tag" style="margin-top:-8px; margin-bottom:12px;">→ console.groq.com (gratuito)</div>', unsafe_allow_html=True)

    tema = st.text_area(
        "TEMA DE INVESTIGACIÓN",
        placeholder="Ej: Impacto del uso de redes sociales en el rendimiento académico de estudiantes universitarios",
        height=90,
    )

    variables_raw = st.text_input(
        "VARIABLES / CATEGORÍAS",
        placeholder="Ej: redes sociales, rendimiento académico",
        help="Separadas por coma"
    )

    col1, col2 = st.columns(2)
    with col1:
        tipo_estudio = st.selectbox("TIPO DE ESTUDIO", ["Cuantitativo", "Cualitativo", "Mixto"])
    with col2:
        tipo_doc = st.selectbox("TIPO DE DOCUMENTO", [
            "Tesis de grado", "Tesis de posgrado", "Artículo científico", "TFM", "Monografía"
        ])

    area = st.selectbox("ÁREA DE CONOCIMIENTO", [
        "Educación / Pedagogía", "Psicología", "Ciencias Sociales",
        "Salud / Medicina", "Ingeniería / Tecnología", "Economía / Administración",
        "Humanidades", "Ciencias Naturales", "Derecho", "Comunicación"
    ])

    col3, col4 = st.columns(2)
    with col3:
        pais = st.text_input("PAÍS", value="Perú", placeholder="Perú")
    with col4:
        norma = st.selectbox("NORMA DE CITA", [
            "APA 7ª edición", "APA 6ª edición", "Vancouver", "Chicago", "MLA", "ISO 690"
        ])

    st.markdown("---")
    generar = st.button("⚡ GENERAR MARCO TEÓRICO")


# ─── MAIN AREA ──────────────────────────────────────────────────────────────

st.markdown("""
<div class="main-header">
  <h1>Generador de Marco Teórico</h1>
  <p>Tesis · Artículos Científicos · TFM &nbsp;|&nbsp; Autores en ES · EN · PT · 中文 &nbsp;|&nbsp; 2021–2026</p>
</div>
""", unsafe_allow_html=True)

# Métricas informativas
st.markdown("""
<div class="metric-row">
  <div class="metric-box"><div class="num">5+</div><div class="lbl">Autores por variable</div></div>
  <div class="metric-box"><div class="num">10</div><div class="lbl">Antecedentes internacionales</div></div>
  <div class="metric-box"><div class="num">8</div><div class="lbl">Antecedentes nacionales</div></div>
  <div class="metric-box"><div class="num">12</div><div class="lbl">Términos en glosario</div></div>
  <div class="metric-box"><div class="num">4</div><div class="lbl">Idiomas de fuentes</div></div>
</div>
""", unsafe_allow_html=True)

st.markdown("""
<div class="info-box">
  📌 <strong>¿Cómo usar?</strong> Completa el formulario en el panel izquierdo, ingresa tu API Key de Groq (gratuita) y presiona <em>Generar Marco Teórico</em>. El proceso tarda entre 3 y 8 minutos según el número de variables.
</div>
""", unsafe_allow_html=True)

# ─── GENERACIÓN ─────────────────────────────────────────────────────────────

if generar:
    # Validaciones
    errores = []
    if not api_key or not api_key.startswith("gsk_"):
        errores.append("⚠️ API Key de Groq inválida (debe empezar con `gsk_`). Obtén la tuya en console.groq.com")
    if not tema.strip():
        errores.append("⚠️ El tema de investigación es obligatorio.")
    if not variables_raw.strip():
        errores.append("⚠️ Ingresa al menos una variable o categoría.")

    if errores:
        for e in errores:
            st.error(e)
        st.stop()

    variables = [v.strip() for v in variables_raw.split(",") if v.strip()]
    datos = {
        "api_key": api_key,
        "tema": tema.strip(),
        "variables": variables,
        "tipo_estudio": tipo_estudio,
        "tipo_doc": tipo_doc,
        "area": area,
        "pais": pais.strip() or "Perú",
        "norma": norma,
    }

    client = get_groq_client(api_key)
    secciones = {}

    total_tasks = len(variables) * 2 + 3
    tasks_done = 0

    # Lista de tareas
    task_labels = []
    for var in variables:
        task_labels.append(f"Definiciones conceptuales: {var}")
        task_labels.append(f"Teorías, modelos y dimensiones: {var}")
    task_labels += [
        "Glosario de términos básicos (12 términos)",
        "Antecedentes internacionales (10 estudios · 2021–2026)",
        f"Antecedentes nacionales — {datos['pais']} (8 estudios · 2021–2026)",
    ]

    st.markdown("### ⚙️ Generando secciones...")
    progress_placeholder = st.empty()
    bar = st.progress(0)
    result_area = st.container()

    def render_progress(current_idx, status="working"):
        html = ""
        for i, label in enumerate(task_labels):
            if i < current_idx:
                icon = '<span class="dot-done">✓</span>'
            elif i == current_idx and status == "working":
                icon = '<span class="dot-working">◉</span>'
            else:
                icon = '<span class="dot-pending">○</span>'
            html += f'<div class="progress-item">{icon} {label}</div>'
        progress_placeholder.markdown(f'<div class="section-card">{html}</div>', unsafe_allow_html=True)

    idx = 0
    for var in variables:
        # Definiciones
        render_progress(idx)
        bar.progress(int(idx / total_tasks * 100))
        secciones[f"def_{var}"] = llamar_groq(client, prompt_definiciones(datos, var))
        idx += 1

        # Teorías
        render_progress(idx)
        bar.progress(int(idx / total_tasks * 100))
        secciones[f"teo_{var}"] = llamar_groq(client, prompt_teorias(datos, var), max_tokens=3200)
        idx += 1
        time.sleep(0.3)

    # Glosario
    render_progress(idx)
    bar.progress(int(idx / total_tasks * 100))
    secciones["glosario"] = llamar_groq(client, prompt_glosario(datos), max_tokens=3000)
    idx += 1

    # Antecedentes int.
    render_progress(idx)
    bar.progress(int(idx / total_tasks * 100))
    secciones["antec_int"] = llamar_groq(client, prompt_antec_internacionales(datos), max_tokens=4000)
    idx += 1
    time.sleep(0.3)

    # Antecedentes nac.
    render_progress(idx)
    bar.progress(int(idx / total_tasks * 100))
    secciones["antec_nac"] = llamar_groq(client, prompt_antec_nacionales(datos), max_tokens=3500)
    idx += 1

    render_progress(idx, status="done")
    bar.progress(100)

    # ── Mostrar resultados ───────────────────────────────────────────────────
    st.success("✅ ¡Marco teórico generado exitosamente!")

    with result_area:
        st.markdown("---")
        st.markdown("## 📋 Vista previa del contenido")

        for var in variables:
            with st.expander(f"📌 Variable: **{var}** — Definiciones conceptuales", expanded=False):
                st.markdown(f'<div class="section-content">{secciones[f"def_{var}"]}</div>', unsafe_allow_html=True)
            with st.expander(f"📐 Variable: **{var}** — Teorías, modelos y dimensiones", expanded=False):
                st.markdown(f'<div class="section-content">{secciones[f"teo_{var}"]}</div>', unsafe_allow_html=True)

        with st.expander("📖 Glosario de términos básicos", expanded=False):
            st.markdown(f'<div class="section-content">{secciones["glosario"]}</div>', unsafe_allow_html=True)

        with st.expander("🌍 Antecedentes internacionales (10 estudios · 2021–2026)", expanded=False):
            st.markdown(f'<div class="section-content">{secciones["antec_int"]}</div>', unsafe_allow_html=True)

        with st.expander(f"🇵🇪 Antecedentes nacionales — {datos['pais']} (8 estudios · 2021–2026)", expanded=False):
            st.markdown(f'<div class="section-content">{secciones["antec_nac"]}</div>', unsafe_allow_html=True)

    # ── Generar y descargar DOCX ─────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### 📥 Descargar documento Word")
    with st.spinner("Construyendo archivo .docx..."):
        docx_bytes = build_docx(datos, secciones)

    slug = tema.strip()[:40].replace(" ", "_").replace("/", "-")
    filename = f"Marco_Teorico_{slug}.docx"

    st.download_button(
        label="⬇️  DESCARGAR MARCO TEÓRICO (.docx)",
        data=docx_bytes,
        file_name=filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    st.markdown(f"""
    <div class="info-box" style="margin-top:1rem;">
      ℹ️ <strong>Recuerda:</strong> Este documento es generado por IA. Verifica las referencias y citas antes de usarlo en tu trabajo académico. El contenido es una base sólida que debes revisar y complementar con fuentes primarias.
    </div>
    """, unsafe_allow_html=True)

else:
    # Estado vacío — instrucciones
    st.markdown("""
    <div class="section-card">
      <div class="section-title">¿Qué genera esta herramienta?</div>
      <div class="section-content">
El marco teórico es uno de los capítulos más exigentes de cualquier investigación académica. Esta herramienta lo genera automáticamente, incluyendo:

<br><br>
<strong>📌 Por cada variable o categoría:</strong>
• Definiciones conceptuales por ≥5 autores (español, inglés, portugués, chino mandarín)
• Teorías y modelos teóricos (mínimo 3) con descripción detallada
• Dimensiones (cuantitativo/mixto) o categorías (cualitativo) con indicadores

<br><br>
<strong>📖 Glosario:</strong>
• 12 términos básicos con definición académica y cita formal

<br><br>
<strong>📚 Antecedentes (todos del período 2021–2026):</strong>
• 10 antecedentes internacionales con resumen/abstract completo
• 8 antecedentes nacionales del país elegido con resumen/abstract completo

<br><br>
<strong>📄 Salida:</strong>
• Documento Word (.docx) descargable con estructura académica profesional
• Citas en el formato de tu elección (APA 7, Vancouver, Chicago, etc.)
      </div>
    </div>
    """, unsafe_allow_html=True)
